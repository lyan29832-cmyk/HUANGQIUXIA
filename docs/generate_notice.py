#!/usr/bin/env python3
"""Generate a temporary quality-system shipment notice matching the existing memo layout."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT_DIR = Path(__file__).resolve().parent

TITLE = "关于俄罗斯客户Nizhpharm Joint Stock Company药品调整发货批次的通知"
DOC_NO = "2026年 第 002 号 T"
ADDRESSEE = "物料部："
BODY = (
    "俄罗斯客户Nizhpharm Joint Stock Company向我司订购药品36盒，"
    "要求剩余效期须大于80%。经确认，正在出库的20251219批次剩余效期"
    "不符合上述要求，无法满足客户效期条件。经与商业部门确认，本次发货"
    "改用20260407批次。该调整与先进先出、近效期先出原则不符。"
    "现根据质量体系要求出具本通知，请按本通知执行出库。"
)
CJK_FONT = "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"


def set_run_font(run, name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)

    style = doc.styles["Normal"]
    style.font.name = "WenQuanYi Micro Hei"
    style.font.size = Pt(16)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "WenQuanYi Micro Hei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run(TITLE)
    set_run_font(run, "WenQuanYi Micro Hei", 18, bold=True)

    number = doc.add_paragraph()
    number.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number.paragraph_format.space_before = Pt(0)
    number.paragraph_format.space_after = Pt(24)
    run = number.add_run(DOC_NO)
    set_run_font(run, "WenQuanYi Micro Hei", 12, bold=False)

    to_p = doc.add_paragraph()
    to_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    to_p.paragraph_format.space_after = Pt(12)
    run = to_p.add_run(ADDRESSEE)
    set_run_font(run, "WenQuanYi Micro Hei", 16, bold=False)

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Cm(0.85)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    body.paragraph_format.line_spacing = 1.75
    body.paragraph_format.space_after = Pt(0)
    run = body.add_run(BODY)
    set_run_font(run, "WenQuanYi Micro Hei", 16, bold=False)

    doc.save(path)


def build_preview_png(path: Path) -> None:
    # A4 at 150 dpi
    width, height = 1240, 1754
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(CJK_FONT, 36)
    meta_font = ImageFont.truetype(CJK_FONT, 24)
    body_font = ImageFont.truetype(CJK_FONT, 30)

    left, right = 140, width - 140
    y = 180

    # Wrap title to at most two lines so it stays centered like the sample.
    title_lines = wrap_text(draw, TITLE, title_font, right - left)
    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        tw = bbox[2] - bbox[0]
        draw.text(((width - tw) / 2, y), line, font=title_font, fill=(20, 20, 20))
        y += 56

    y += 28
    bbox = draw.textbbox((0, 0), DOC_NO, font=meta_font)
    tw = bbox[2] - bbox[0]
    draw.text((right - tw, y), DOC_NO, font=meta_font, fill=(30, 30, 30))

    y += 70
    draw.text((left, y), ADDRESSEE, font=body_font, fill=(20, 20, 20))

    y += 58
    indent = 60
    max_width = right - left
    lines = wrap_text(draw, BODY, body_font, max_width)
    for i, line in enumerate(lines):
        x = left + indent if i == 0 else left
        draw.text((x, y), line, font=body_font, fill=(20, 20, 20))
        y += 52

    img.save(path, "PNG")


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx(OUT_DIR / "关于俄罗斯客户Nizhpharm药品调整发货批次的通知.docx")
    build_preview_png(OUT_DIR / "关于俄罗斯客户Nizhpharm药品调整发货批次的通知.png")


if __name__ == "__main__":
    main()
